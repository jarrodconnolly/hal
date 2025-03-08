import os
from pathlib import Path
import pypdfium2 as pdfium
from sentence_transformers import SentenceTransformer
import numpy as np
import torch
import time
from multiprocessing import Pool
from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document as DocxDocument
import json
from qdrant_client import QdrantClient
from qdrant_client.http import models as rest

DATA_DIR = os.path.expanduser("~/data")
OUTPUT_DIR = "vector_db"
CHUNK_SIZE = 2000
EMBEDDING_MODEL = "thenlper/gte-large"
HNSW_M = 32

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Using device: {device}")

embedding_model = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

def extract_text_from_pdf(file_path):
    try:
        pdf = pdfium.PdfDocument(file_path)
        text = ""
        for i in range(len(pdf)):
            text += pdf[i].get_textpage().get_text_bounded() or ""
        return text
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def extract_text_from_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

def split_text(text, chunk_size=CHUNK_SIZE):
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

def process_file(file_path):
    if file_path.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith((".txt", ".md")):
        text = extract_text_from_txt(file_path)
    elif file_path.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        return [], []
    if text:
        file_chunks = split_text(text)
        return file_chunks, [(file_path, i) for i in range(len(file_chunks))]
    return [], []

def process_files(file_paths):
    start_time = time.time()
    chunks = []
    metadata = []
    with Pool(processes=20) as pool:
        results = pool.map(process_file, file_paths)
    for file_chunks, file_metadata in results:
        if file_chunks:
            chunks.extend(file_chunks)
            metadata.extend(file_metadata)
            print(f"Processed {os.path.basename(file_metadata[0][0])}: {len(file_chunks)} chunks")
    elapsed = time.time() - start_time
    print(f"Text extraction and chunking took {elapsed:.2f} seconds")
    return chunks, metadata

def generate_embeddings(chunks, model_name=EMBEDDING_MODEL):
    start_time = time.time()
    model = SentenceTransformer(model_name, device=device)
    embeddings = model.encode(chunks, batch_size=256, show_progress_bar=True)
    elapsed = time.time() - start_time
    print(f"Embedding generation took {elapsed:.2f} seconds")
    return embeddings

def update_vector_store(data_dir, output_dir):
    total_start = time.time()
    state_file = os.path.join(output_dir, "state.json")
    
    prev_state = {}
    if os.path.exists(state_file):
        with open(state_file, "r") as f:
            prev_state = json.load(f)

    current_state = {}
    file_paths = [os.path.join(root, file) for root, _, files in os.walk(data_dir) for file in files]
    for fp in file_paths:
        if fp.lower().endswith((".pdf", ".txt", ".md", ".docx")):
            current_state[fp] = os.path.getmtime(fp)

    client = QdrantClient("localhost", port=6333, timeout=300)
    collection_name = "hal_docs"

    # Reset collection if state mismatch (full rebuild), but skip delete on first run
    if prev_state and current_state != prev_state:  # Only delete if prev_state exists
        if client.collection_exists(collection_name):
            print("Resetting collection for full rebuild.")
            try:
                client.delete_collection(collection_name)
            except Exception as e:
                print(f"Warning: Failed to delete collection: {str(e)} - Proceeding with create.")

    if not client.collection_exists(collection_name):
        print("Creating new collection.")
        client.create_collection(
            collection_name=collection_name,
            vectors_config=rest.VectorParams(
                size=1024,
                distance=rest.Distance.COSINE
            ),
            hnsw_config=rest.HnswConfigDiff(
                m=HNSW_M
            )
        )

    if not file_paths:
        print("No files to process.")
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        with open(state_file, "w") as f:
            json.dump(current_state, f)
        return

    to_process = [fp for fp in file_paths if fp not in prev_state or prev_state[fp] != current_state[fp]]
    if not to_process:
        print("No changes detected.")
        return

    new_chunks, new_metadata = process_files(to_process)
    if not new_chunks:
        print("No new/changed data to process.")
        return

    # Generate embeddings for new chunks
    embeddings = generate_embeddings(new_chunks)

    # Get current point count for unique IDs
    collection_info = client.get_collection(collection_name)
    next_id = collection_info.points_count

    # Validate and prepare points
    points = []
    for idx, (embedding, m, chunk) in enumerate(zip(embeddings, new_metadata, new_chunks)):
        if len(embedding) != 1024:
            print(f"Warning: Embedding {idx} has {len(embedding)} dims, expected 1024")
        points.append(
            rest.PointStruct(
                id=next_id + idx,
                vector=embedding.tolist(),
                payload={"source": m[0], "chunk_id": m[1], "content": chunk}
            )
        )

    # Upsert in batches with error handling
    batch_size = 1000
    for i in range(0, len(points), batch_size):
        batch = points[i:i + batch_size]
        print(f"Upserting batch {i // batch_size + 1} of {(len(points) - 1) // batch_size + 1} ({len(batch)} points)")
        try:
            client.upsert(collection_name=collection_name, points=batch)
        except Exception as e:
            print(f"Upsert failed: {str(e)}")
            if hasattr(e, "content"):
                print(f"Response content: {e.content}")
            raise
    
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    with open(state_file, "w") as f:
        json.dump(current_state, f)
    
    total_chunks = next_id + len(new_chunks)
    print(f"Updated with {len(new_chunks)} new/changed chunks, total {total_chunks} in {time.time() - total_start:.2f} seconds")

def main():
    print("Starting text extraction, embedding, and storage process...")
    update_vector_store(DATA_DIR, OUTPUT_DIR)
    print("Process complete!")

if __name__ == "__main__":
    main()